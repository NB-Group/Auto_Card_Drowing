"""
春秋杀卡牌数据导出Word文档脚本
将cards.json中的卡牌数据转换为格式化的Word文档，方便汇报使用
"""

import json
import os
import glob
from datetime import datetime
from collections import defaultdict

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("❌ 缺少python-docx库，请先安装：")
    print("   pip install python-docx")
    exit(1)

class CardsToWordExporter:
    def __init__(self):
        self.base_path = os.path.dirname(os.path.abspath(__file__))
        self.cards_file = os.path.join(self.base_path, "cards.json")
        self.output_file = os.path.join(self.base_path, f"春秋杀卡牌汇总_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
        
        # 卡牌类型emoji映射
        self.type_emojis = {
            "国家卡": "🏰",
            "思想卡": "🧠", 
            "变法卡": "⚖️",
            "连锁卡": "🔗",
            "军事卡": "⚔️",
            "经济卡": "💰",
            "道具卡": "🎁",
            "锦囊牌": "📜",
            "祭祀卡": "🙏"
        }
    
    def load_cards_data(self):
        """加载卡牌数据"""
        try:
            with open(self.cards_file, 'r', encoding='utf-8') as f:
                cards_data = json.load(f)
                print(f"✅ 成功加载 {len(cards_data)} 张卡牌数据")
                return cards_data
        except Exception as e:
            print(f"❌ 读取cards.json失败: {e}")
            return []
    
    def group_cards_by_type(self, cards_data):
        """按卡牌类型分组"""
        grouped_cards = defaultdict(list)
        for card in cards_data:
            card_group = card.get('card_group', '未分类')
            grouped_cards[card_group].append(card)
        return dict(grouped_cards)
    
    def add_title_page(self, doc):
        """添加标题页"""
        # 主标题
        title = doc.add_heading('春秋杀卡牌设计汇总', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 副标题
        subtitle = doc.add_paragraph('Spring-Autumn Kill Card Game Design Summary')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(14)
        subtitle.runs[0].font.name = '微软雅黑'
        
        # 添加空行
        doc.add_paragraph()
        
        # 项目信息
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = info_para.add_run(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}\n")
        run2 = info_para.add_run("项目：春秋杀桌游卡牌生成器\n")
        run3 = info_para.add_run("技术：AI生成 + 自动化合成")
        
        # 设置字体
        for run in [run1, run2, run3]:
            run.font.name = '微软雅黑'
            run.font.size = Pt(12)
        
        # 分页
        doc.add_page_break()
    
    def add_statistics_section(self, doc, grouped_cards):
        """添加统计信息部分"""
        doc.add_heading('📊 卡牌统计概览', level=1)
        
        # 总体统计
        total_cards = sum(len(cards) for cards in grouped_cards.values())
        stats_para = doc.add_paragraph()
        run1 = stats_para.add_run(f"📋 总卡牌数量：{total_cards} 张\n")
        run2 = stats_para.add_run(f"🏷️ 卡牌类型：{len(grouped_cards)} 种\n")
        run3 = stats_para.add_run(f"🎯 设计理念：基于春秋战国历史背景的策略卡牌游戏")
        
        # 设置字体
        for run in [run1, run2, run3]:
            run.font.name = '微软雅黑'
            run.font.size = Pt(11)
        
        # 添加统计表格
        table_desc_para = doc.add_paragraph()
        table_desc_run = table_desc_para.add_run("\n📈 各类型卡牌数量分布：")
        table_desc_run.font.name = '微软雅黑'
        table_desc_run.font.size = Pt(11)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 表头
        header_cells = table.rows[0].cells
        header_cells[0].text = '卡牌类型'
        header_cells[1].text = 'emoji'
        header_cells[2].text = '数量'
        header_cells[3].text = '占比'
        
        # 设置表头格式
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.name = '微软雅黑'
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加数据行
        for card_type, cards in grouped_cards.items():
            row_cells = table.add_row().cells
            row_cells[0].text = card_type
            row_cells[1].text = self.type_emojis.get(card_type, '')
            row_cells[2].text = str(len(cards))
            row_cells[3].text = f"{len(cards)/total_cards*100:.1f}%"
            
            # 居中对齐并设置字体
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in cell.paragraphs[0].runs:
                    run.font.name = '微软雅黑'
        
        # 统一设置整个表格字体
        self.set_table_font(table)
        
        doc.add_paragraph()
    
    def add_game_rules_section(self, doc):
        """添加游戏规则部分"""
        doc.add_heading('🎮 游戏规则概述', level=1)
        
        rules_text = """
春秋杀是一款基于春秋战国历史背景的策略卡牌游戏，玩家扮演各国君主，通过变法、军事、经济等手段争夺霸权。

🏛️ 核心机制：
• 国家属性：军事、经济、政治三大属性系统
• 地形加成：河流沿岸、山地等地形提供不同加成
• 思想切换：五种思想流派，提供不同属性加成
• 连锁反应：变法卡触发连锁卡，增加策略深度

⚔️ 游戏阶段：
1. 朝堂议事阶段：出变法卡和连锁卡
2. 农耕阶段：抽取军事和经济卡牌
3. 军事阶段：进行战争和军事行动
4. 会盟阶段：外交协商和结盟
5. 祭祀阶段：随机事件和BUFF效果

🎯 胜利条件：
通过军事征服、经济发展、政治影响等多种途径获得胜利点数。
        """
        
        rules_para = doc.add_paragraph()
        rules_run = rules_para.add_run(rules_text.strip())
        rules_run.font.name = '微软雅黑'
        rules_run.font.size = Pt(11)
        doc.add_paragraph()
    
    def add_cards_detail_section(self, doc, grouped_cards):
        """添加卡牌详细信息部分"""
        doc.add_heading('🃏 卡牌详细信息', level=1)
        
        for card_type, cards in grouped_cards.items():
            # 添加卡牌类型标题
            emoji = self.type_emojis.get(card_type, '')
            type_heading = doc.add_heading(f'{emoji} {card_type} ({len(cards)}张)', level=2)
            
            # 创建卡牌信息表格
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light List Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # 表头
            header_cells = table.rows[0].cells
            header_cells[0].text = '卡牌名称'
            header_cells[1].text = '价格'
            header_cells[2].text = '主题色'
            header_cells[3].text = 'AI提示词'
            header_cells[4].text = '效果描述'
            
            # 设置表头格式
            for cell in header_cells:
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.name = '微软雅黑'
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加卡牌数据
            for card in cards:
                row_cells = table.add_row().cells
                row_cells[0].text = card.get('card_name', '')
                row_cells[1].text = card.get('price', '')
                row_cells[2].text = card.get('color_theme', '')
                row_cells[3].text = card.get('ai_prompt', '')
                row_cells[4].text = card.get('description', '')
                
                # 设置单元格格式
                row_cells[0].paragraphs[0].runs[0].font.bold = True
                row_cells[0].paragraphs[0].runs[0].font.name = '微软雅黑'
                row_cells[1].paragraphs[0].runs[0].font.bold = True  # 价格也加粗
                row_cells[1].paragraphs[0].runs[0].font.name = '微软雅黑'
                row_cells[3].width = Inches(2.8)  # AI提示词列宽一些
                row_cells[4].width = Inches(2.2)  # 效果描述列宽一些
            
            # 统一设置整个表格字体
            self.set_table_font(table)
            
            # 设置表格列宽
            for row in table.rows:
                row.cells[0].width = Inches(1.2)  # 卡牌名称
                row.cells[1].width = Inches(0.8)  # 价格
                row.cells[2].width = Inches(0.8)  # 主题色
                row.cells[3].width = Inches(2.8)  # AI提示词
                row.cells[4].width = Inches(2.2)  # 效果描述
            
            doc.add_paragraph()  # 添加间距
    
    def add_technical_section(self, doc):
        """添加技术实现部分"""
        doc.add_heading('⚙️ 技术实现', level=1)
        
        tech_text = """
🤖 AI图片生成：
• 使用Microsoft Copilot AI生成卡牌插画
• 国风写实融合风格，低饱和度复古色调
• 自动化浏览器操作，支持批量生成

🎨 卡牌合成系统：
• Python + Pillow图像处理库
• 智能文字布局和边距控制
• 渐变融合效果，消除图片割裂感
• 主题色智能emoji渲染

🛠️ 核心技术栈：
• Python 3.8+ - 主要编程语言
• Playwright - 浏览器自动化
• Pillow (PIL) - 图像处理
• JSON - 卡牌数据配置

✨ 用户体验优化：
• 彩色日志输出系统
• 动态进度条显示
• 智能错误处理和重试
• 浏览器状态保持

📊 输出格式：
• PNG格式高质量卡牌图片
• 标准卡牌尺寸适配
• 支持批量导出和单卡生成
        """
        
        tech_para = doc.add_paragraph()
        tech_run = tech_para.add_run(tech_text.strip())
        tech_run.font.name = '微软雅黑'
        tech_run.font.size = Pt(11)
    
    def set_font_to_microsoft_yahei(self, element, size=Pt(11), bold=False):
        """统一设置字体为微软雅黑"""
        if hasattr(element, 'runs'):
            for run in element.runs:
                run.font.name = '微软雅黑'
                run.font.size = size
                run.font.bold = bold
        elif hasattr(element, 'font'):
            element.font.name = '微软雅黑'
            element.font.size = size
            element.font.bold = bold
    
    def set_table_font(self, table):
        """设置表格字体为微软雅黑"""
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    self.set_font_to_microsoft_yahei(paragraph, size=Pt(10))
    
    def set_document_style(self, doc):
        """设置文档样式"""
        # 设置默认字体
        style = doc.styles['Normal']
        font = style.font
        font.name = '微软雅黑'
        font.size = Pt(11)
        
        # 设置所有标题样式为微软雅黑
        for i in range(10):  # 设置标题1-9级
            try:
                heading_style = doc.styles[f'Heading {i}']
                heading_style.font.name = '微软雅黑'
                heading_style.font.bold = True
                if i == 0:  # 主标题
                    heading_style.font.size = Pt(20)
                elif i == 1:  # 一级标题
                    heading_style.font.size = Pt(16)
                elif i == 2:  # 二级标题
                    heading_style.font.size = Pt(14)
                else:
                    heading_style.font.size = Pt(12)
            except:
                pass  # 忽略不存在的标题样式
        
        # 设置表格默认样式
        try:
            table_style = doc.styles['Table Grid']
            table_style.font.name = '微软雅黑'
            table_style.font.size = Pt(10)
        except:
            pass
    
    def export_to_word(self):
        """导出到Word文档"""
        print("🚀 开始生成Word文档...")
        
        # 加载卡牌数据
        cards_data = self.load_cards_data()
        if not cards_data:
            return False
        
        # 按类型分组
        grouped_cards = self.group_cards_by_type(cards_data)
        
        # 创建Word文档
        doc = Document()
        self.set_document_style(doc)
        
        # 添加各个部分
        print("📝 添加标题页...")
        self.add_title_page(doc)
        
        print("📊 添加统计信息...")
        self.add_statistics_section(doc, grouped_cards)
        
        print("🎮 添加游戏规则...")
        self.add_game_rules_section(doc)
        
        print("🃏 添加卡牌详细信息...")
        self.add_cards_detail_section(doc, grouped_cards)
        
        print("⚙️ 添加技术实现...")
        self.add_technical_section(doc)
        
        # 保存文档
        try:
            doc.save(self.output_file)
            print(f"✅ Word文档生成成功！")
            print(f"📄 文件路径：{self.output_file}")
            return True
        except Exception as e:
            print(f"❌ 保存Word文档失败：{e}")
            return False

def main():
    """主函数"""
    print("=" * 50)
    print("🌟 春秋杀卡牌数据导出工具")
    print("=" * 50)
    
    # 🗑️ 删除旧的卡牌汇总文件
    old_files = glob.glob("春秋杀卡牌汇总_*.docx")
    if old_files:
        print("🗑️ 清理旧的卡牌汇总文件...")
        for file in old_files:
            try:
                os.remove(file)
                print(f"   ✅ 已删除：{file}")
            except:
                print(f"   ❌ 删除失败：{file}")
        print()
    
    exporter = CardsToWordExporter()
    success = exporter.export_to_word()
    
    if success:
        print("\n🎉 导出完成！可以用于汇报展示了！")
    else:
        print("\n💥 导出失败，请检查错误信息")

if __name__ == "__main__":
    main() 