#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
春秋杀项目 - 综合汇报生成器
生成包含所有功能、数据、分析的Word汇报文档
"""

import json
import os
import subprocess
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import glob

class ComprehensiveReport:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        
    def setup_styles(self):
        """设置文档样式"""
        # 标题样式
        title_style = self.doc.styles['Heading 1']
        title_font = title_style.font
        title_font.name = '微软雅黑'
        title_font.size = Pt(18)
        title_font.color.rgb = RGBColor(47, 85, 151)
        
        # 二级标题样式
        subtitle_style = self.doc.styles['Heading 2']
        subtitle_font = subtitle_style.font
        subtitle_font.name = '微软雅黑'
        subtitle_font.size = Pt(14)
        subtitle_font.color.rgb = RGBColor(68, 114, 196)
        
        # 正文样式
        normal_style = self.doc.styles['Normal']
        normal_font = normal_style.font
        normal_font.name = '微软雅黑'
        normal_font.size = Pt(10)
    
    def add_title_page(self):
        """添加封面页"""
        # 主标题
        title = self.doc.add_heading('🏰 春秋杀卡牌游戏项目', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = self.doc.add_heading('📊 综合功能汇报文档', 2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 项目信息
        info_paragraph = self.doc.add_paragraph()
        info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_paragraph.add_run(f"\n\n📅 生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}\n")
        info_paragraph.add_run("🎯 项目版本：v2.0\n")
        info_paragraph.add_run("🎮 游戏类型：策略卡牌游戏\n")
        info_paragraph.add_run("🎨 技术栈：Python + AI图像生成 + Excel数据管理\n")
        
        # 分页
        self.doc.add_page_break()
    
    def add_card_summary_only(self):
        """添加纯卡牌功能汇总"""
        self.doc.add_heading('🃏 卡牌功能汇总', 1)
        
        # 读取卡牌数据
        try:
            with open('cards.json', 'r', encoding='utf-8') as f:
                cards_data = json.load(f)
        except:
            cards_data = []
        
        if not cards_data:
            self.doc.add_paragraph("❌ 未找到卡牌数据文件")
            return
            
        # 军事卡牌汇总
        self.doc.add_heading('⚔️ 军事卡牌', 2)
        military_cards = [card for card in cards_data if card.get('card_group') == '军事卡']
        
        military_table = self.doc.add_table(rows=1, cols=2)
        military_table.style = 'Light Grid Accent 1'
        
        header_cells = military_table.rows[0].cells
        header_cells[0].text = '🃏 卡牌名称'
        header_cells[1].text = '📝 卡牌描述'
        
        for card in military_cards:
            row_cells = military_table.add_row().cells
            row_cells[0].text = card.get('card_name', '')
            row_cells[1].text = card.get('description', '')
        
        # 经济卡牌汇总
        self.doc.add_heading('💰 经济卡牌', 2)
        economy_cards = [card for card in cards_data if card.get('card_group') == '经济卡']
        
        economy_table = self.doc.add_table(rows=1, cols=2)
        economy_table.style = 'Light Grid Accent 1'
        
        header_cells = economy_table.rows[0].cells
        header_cells[0].text = '🃏 卡牌名称'
        header_cells[1].text = '📝 卡牌描述'
        
        for card in economy_cards:
            row_cells = economy_table.add_row().cells
            row_cells[0].text = card.get('card_name', '')
            row_cells[1].text = card.get('description', '')
        
        # 锦囊牌汇总（单独处理，确保显示）
        self.doc.add_heading('📜 锦囊牌', 2)
        jinlang_cards = [card for card in cards_data if card.get('card_group') == '锦囊牌']
        
        jinlang_table = self.doc.add_table(rows=1, cols=2)
        jinlang_table.style = 'Light Grid Accent 1'
        
        header_cells = jinlang_table.rows[0].cells
        header_cells[0].text = '🃏 卡牌名称'
        header_cells[1].text = '📝 卡牌描述'
        
        for card in jinlang_cards:
            row_cells = jinlang_table.add_row().cells
            row_cells[0].text = card.get('card_name', '')
            row_cells[1].text = card.get('description', '')
        
        # 其他卡牌类型汇总（排除军事卡、经济卡、锦囊牌）
        other_groups = set(card.get('card_group') for card in cards_data) - {'军事卡', '经济卡', '锦囊牌'}
        
        for group in sorted(other_groups):
            if not group:
                continue
                
            group_cards = [card for card in cards_data if card.get('card_group') == group]
            if group_cards:
                self.doc.add_heading(f'🎯 {group}', 2)
                
                group_table = self.doc.add_table(rows=1, cols=2)
                group_table.style = 'Light Grid Accent 1'
                
                header_cells = group_table.rows[0].cells
                header_cells[0].text = '🃏 卡牌名称'
                header_cells[1].text = '📝 卡牌描述'
                
                for card in group_cards:
                    row_cells = group_table.add_row().cells
                    row_cells[0].text = card.get('card_name', '')
                    row_cells[1].text = card.get('description', '')
        
        # 卡牌功能统计
        self.doc.add_heading('📊 卡牌统计概览', 2)
        stats_text = f"""
🃏 卡牌总数：{len(cards_data)} 张

📋 分类统计：
• ⚔️ 军事卡牌：{len([c for c in cards_data if c.get('card_group') == '军事卡'])} 张
• 💰 经济卡牌：{len([c for c in cards_data if c.get('card_group') == '经济卡'])} 张
• 🏰 国家卡牌：{len([c for c in cards_data if c.get('card_group') == '国家卡'])} 张
• 🧠 思想卡牌：{len([c for c in cards_data if c.get('card_group') == '思想卡'])} 张
• ⚖️ 变法卡牌：{len([c for c in cards_data if c.get('card_group') == '变法卡'])} 张
• 🔗 连锁卡牌：{len([c for c in cards_data if c.get('card_group') == '连锁卡'])} 张
• 🎁 道具卡牌：{len([c for c in cards_data if c.get('card_group') == '道具卡'])} 张
• 📜 锦囊卡牌：{len([c for c in cards_data if c.get('card_group') == '锦囊牌'])} 张
• 🙏 祭祀卡牌：{len([c for c in cards_data if c.get('card_group') == '祭祀卡'])} 张
        """
        self.doc.add_paragraph(stats_text)
    
    def generate_report(self, filename=None):
        """生成综合汇报"""
        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"春秋杀项目综合汇报_{timestamp}.docx"
        
        print("🚀 开始生成综合汇报...")
        
        self.add_card_summary_only()
        print("✅ 卡牌功能汇总完成")
        
        # 保存文档
        self.doc.save(filename)
        print(f"🎉 综合汇报生成完成：{filename}")
        return filename

def main():
    """主函数"""
    print("📊 春秋杀项目综合汇报生成器")
    print("=" * 50)
    
    # 🗑️ 删除旧的综合汇报文件和游戏记录表
    old_reports = glob.glob("春秋杀项目综合汇报_*.docx")
    old_sheets = glob.glob("春秋杀游戏记录表_*.xlsx")
    old_files = old_reports + old_sheets
    
    if old_files:
        print("🗑️ 清理旧文件...")
        for file in old_files:
            try:
                os.remove(file)
                print(f"   ✅ 已删除：{file}")
            except:
                print(f"   ❌ 删除失败：{file}")
        print()
    
    # 检查是否有最新的游戏记录表
    try:
        import country_data_sheet
        print("📈 正在生成最新的游戏记录表...")
        sheet_generator = country_data_sheet.GameRecordSheet()
        excel_file = sheet_generator.generate_sheet()
        print(f"✅ 游戏记录表生成完成：{excel_file}")
    except Exception as e:
        print(f"⚠️  游戏记录表生成失败：{e}")
    
    # 生成综合汇报
    report_generator = ComprehensiveReport()
    report_file = report_generator.generate_report()
    
    print("\n📋 汇报内容包含：")
    print("   • 🃏 卡牌功能汇总")
    
    print(f"\n🎯 汇报文档：{report_file}")
    print("📊 可直接用于项目汇报展示！")

if __name__ == "__main__":
    main() 